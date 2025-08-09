import json
import os
import re
from typing import Dict, List, Optional, Union
from openai import OpenAI
import base64
from pathlib import Path
import tempfile

# Import các thư viện xử lý document
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("Warning: python-docx not installed. DOCX processing will be limited.")

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    print("Warning: PDF processing libraries not installed.")

try:
    from PIL import Image
    import pytesseract
except ImportError:
    print("Warning: OCR libraries not installed.")

# Khởi tạo OpenAI client
client = OpenAI(
    base_url="https://aiportalapi.stu-platform.live/jpe",
    api_key="sk-xvI5gYSbiDQ3c4blptwn0A"
)

class DocumentExtractor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png']
        self.form_templates_file = "./database/form_templates.json"
        self.ensure_form_templates()
        self.load_form_templates()
    
    def ensure_form_templates(self):
        """Đảm bảo file templates biểu mẫu tồn tại"""
        os.makedirs(os.path.dirname(self.form_templates_file), exist_ok=True)
        if not os.path.exists(self.form_templates_file):
            default_templates = {
                "employee_info_form": {
                    "name": "Thông tin nhân viên",
                    "fields": [
                        {"field": "full_name", "label": "Họ và tên", "type": "text", "required": True},
                        {"field": "email", "label": "Email", "type": "email", "required": True},
                        {"field": "phone", "label": "Số điện thoại", "type": "text", "required": True},
                        {"field": "address", "label": "Địa chỉ", "type": "text", "required": False},
                        {"field": "birth_date", "label": "Ngày sinh", "type": "date", "required": False},
                        {"field": "id_number", "label": "CMND/CCCD", "type": "text", "required": True},
                        {"field": "education", "label": "Trình độ học vấn", "type": "text", "required": False},
                        {"field": "experience", "label": "Kinh nghiệm làm việc", "type": "text", "required": False},
                        {"field": "skills", "label": "Kỹ năng", "type": "text", "required": False},
                        {"field": "position_applied", "label": "Vị trí ứng tuyển", "type": "text", "required": True}
                    ]
                },
                "contract_form": {
                    "name": "Thông tin hợp đồng",
                    "fields": [
                        {"field": "employee_name", "label": "Tên nhân viên", "type": "text", "required": True},
                        {"field": "position", "label": "Chức vụ", "type": "text", "required": True},
                        {"field": "department", "label": "Phòng ban", "type": "text", "required": True},
                        {"field": "start_date", "label": "Ngày bắt đầu", "type": "date", "required": True},
                        {"field": "salary", "label": "Mức lương", "type": "number", "required": True},
                        {"field": "contract_type", "label": "Loại hợp đồng", "type": "text", "required": True},
                        {"field": "probation_period", "label": "Thời gian thử việc", "type": "text", "required": False}
                    ]
                }
            }
            
            with open(self.form_templates_file, 'w', encoding='utf-8') as f:
                json.dump(default_templates, f, ensure_ascii=False, indent=2)
    
    def load_form_templates(self):
        """Tải templates biểu mẫu từ file"""
        with open(self.form_templates_file, 'r', encoding='utf-8') as f:
            self.form_templates = json.load(f)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Trích xuất text từ file PDF"""
        text = ""
        try:
            # Thử với pdfplumber trước (tốt hơn cho layout)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # Nếu không có text, thử với PyPDF2
            if not text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                        
        except Exception as e:
            print(f"Error extracting PDF: {e}")
            
        return text.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Trích xuất text từ file DOCX"""
        text = ""
        try:
            doc = Document(file_path)
            
            # Trích xuất text từ paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Trích xuất text từ tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
                    
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            
        return text.strip()
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Trích xuất text từ ảnh bằng OCR"""
        text = ""
        try:
            image = Image.open(file_path)
            # Cải thiện chất lượng ảnh cho OCR
            image = image.convert('RGB')
            text = pytesseract.image_to_string(image, lang='vie+eng')
        except Exception as e:
            print(f"Error extracting text from image: {e}")
            
        return text.strip()
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Trích xuất text từ file dựa trên extension"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            return self.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def extract_cv_information(self, text: str) -> Dict:
        """Trích xuất thông tin từ CV bằng AI"""
        
        prompt = f"""
Từ nội dung CV sau, hãy trích xuất thông tin có cấu trúc:

{text}

Trích xuất các thông tin sau (nếu có):
1. Thông tin cá nhân: Họ tên, email, số điện thoại, địa chỉ, ngày sinh
2. Học vấn: Trường, chuyên ngành, năm tốt nghiệp, bằng cấp
3. Kinh nghiệm làm việc: Công ty, vị trí, thời gian, mô tả công việc
4. Kỹ năng: Kỹ năng chuyên môn, kỹ năng mềm, ngôn ngữ
5. Chứng chỉ: Tên chứng chỉ, tổ chức cấp, năm cấp
6. Dự án: Tên dự án, vai trò, công nghệ sử dụng
7. Sở thích: Hoạt động ngoại khóa, sở thích cá nhân

Trả về format JSON với cấu trúc rõ ràng. Nếu không tìm thấy thông tin nào thì để giá trị null.

```json
{{
  "personal_info": {{
    "full_name": "string",
    "email": "string", 
    "phone": "string",
    "address": "string",
    "birth_date": "string"
  }},
  "education": [
    {{
      "school": "string",
      "major": "string",
      "degree": "string",
      "graduation_year": "string"
    }}
  ],
  "experience": [
    {{
      "company": "string",
      "position": "string",
      "duration": "string",
      "description": "string"
    }}
  ],
  "skills": {{
    "technical": ["string"],
    "soft_skills": ["string"],
    "languages": ["string"]
  }},
  "certifications": [
    {{
      "name": "string",
      "issuer": "string",
      "year": "string"
    }}
  ],
  "projects": [
    {{
      "name": "string",
      "role": "string",
      "technologies": ["string"],
      "description": "string"
    }}
  ],
  "interests": ["string"]
}}
```
"""
        
        try:
            response = client.chat.completions.create(
                model="GPT-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.3
            )
            
            content = response.choices[0].message.content
            
            # Parse JSON từ response
            json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
            if json_match:
                try:
                    cv_data = json.loads(json_match.group(1))
                    return cv_data
                except json.JSONDecodeError:
                    pass
            
            # Thử parse toàn bộ response
            try:
                cv_data = json.loads(content)
                return cv_data
            except json.JSONDecodeError:
                return {"error": "Could not parse CV information", "raw_content": content}
                
        except Exception as e:
            return {"error": f"Failed to extract CV information: {e}"}
    
    def extract_document_information(self, text: str, document_type: str = "general") -> Dict:
        """Trích xuất thông tin từ giấy tờ khác"""
        
        if document_type == "id_card":
            prompt = f"""
Từ nội dung CMND/CCCD sau, trích xuất thông tin:

{text}

Trích xuất:
- Số CMND/CCCD
- Họ và tên
- Ngày sinh
- Giới tính
- Quê quán
- Thường trú
- Ngày cấp
- Nơi cấp

Format JSON:
```json
{{
  "id_number": "string",
  "full_name": "string",
  "birth_date": "string",
  "gender": "string",
  "hometown": "string",
  "permanent_address": "string",
  "issue_date": "string",
  "issue_place": "string"
}}
```
"""
        elif document_type == "diploma":
            prompt = f"""
Từ nội dung bằng cấp sau, trích xuất thông tin:

{text}

Trích xuất:
- Tên người nhận
- Loại bằng
- Chuyên ngành
- Trường/Tổ chức cấp
- Năm tốt nghiệp
- Xếp loại

Format JSON:
```json
{{
  "recipient_name": "string",
  "degree_type": "string",
  "major": "string",
  "institution": "string",
  "graduation_year": "string",
  "classification": "string"
}}
```
"""
        else:  # general document
            prompt = f"""
Từ nội dung tài liệu sau, trích xuất các thông tin quan trọng:

{text}

Tìm và trích xuất:
- Tên tài liệu
- Ngày tháng
- Tên người/tổ chức liên quan
- Số hiệu tài liệu
- Thông tin quan trọng khác

Format JSON với các field phù hợp.
"""
        
        try:
            response = client.chat.completions.create(
                model="GPT-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0.3
            )
            
            content = response.choices[0].message.content
            
            # Parse JSON
            json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
            if json_match:
                try:
                    doc_data = json.loads(json_match.group(1))
                    return doc_data
                except json.JSONDecodeError:
                    pass
            
            try:
                doc_data = json.loads(content)
                return doc_data
            except json.JSONDecodeError:
                return {"error": "Could not parse document information", "raw_content": content}
                
        except Exception as e:
            return {"error": f"Failed to extract document information: {e}"}
    
    def auto_fill_form(self, extracted_data: Dict, form_template: str) -> Dict:
        """Tự động điền biểu mẫu từ dữ liệu đã trích xuất"""
        
        if form_template not in self.form_templates:
            return {"error": f"Form template '{form_template}' not found"}
        
        template = self.form_templates[form_template]
        filled_form = {
            "form_name": template["name"],
            "fields": {},
            "missing_fields": [],
            "confidence": {}
        }
        
        # Mapping logic với AI
        prompt = f"""
Dựa trên dữ liệu đã trích xuất và template biểu mẫu, hãy điền thông tin vào các trường:

Dữ liệu có sẵn:
{json.dumps(extracted_data, ensure_ascii=False, indent=2)}

Template biểu mẫu:
{json.dumps(template["fields"], ensure_ascii=False, indent=2)}

Hãy map dữ liệu vào các trường phù hợp và đánh giá độ tin cậy (0-1).

Format trả về:
```json
{{
  "field_name": {{
    "value": "giá trị điền vào",
    "confidence": 0.95,
    "source": "trường nào trong dữ liệu gốc"
  }}
}}
```

Chỉ điền những trường có thông tin rõ ràng.
"""
        
        try:
            response = client.chat.completions.create(
                model="GPT-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
                temperature=0.2
            )
            
            content = response.choices[0].message.content
            
            # Parse mapping result
            json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
            if json_match:
                try:
                    mapping_result = json.loads(json_match.group(1))
                    
                    # Xử lý kết quả mapping
                    for field_info in template["fields"]:
                        field_name = field_info["field"]
                        
                        if field_name in mapping_result:
                            filled_form["fields"][field_name] = mapping_result[field_name]["value"]
                            filled_form["confidence"][field_name] = mapping_result[field_name]["confidence"]
                        else:
                            if field_info["required"]:
                                filled_form["missing_fields"].append(field_name)
                    
                    return filled_form
                    
                except json.JSONDecodeError:
                    pass
            
            # Fallback: simple mapping
            return self.simple_form_mapping(extracted_data, template)
            
        except Exception as e:
            return {"error": f"Failed to auto-fill form: {e}"}
    
    def simple_form_mapping(self, extracted_data: Dict, template: Dict) -> Dict:
        """Mapping đơn giản khi AI không hoạt động"""
        filled_form = {
            "form_name": template["name"],
            "fields": {},
            "missing_fields": [],
            "confidence": {}
        }
        
        # Simple mapping rules
        mapping_rules = {
            "full_name": ["personal_info.full_name", "recipient_name", "full_name"],
            "email": ["personal_info.email", "email"],
            "phone": ["personal_info.phone", "phone"],
            "address": ["personal_info.address", "permanent_address", "address"],
            "birth_date": ["personal_info.birth_date", "birth_date"],
            "id_number": ["id_number"],
            "education": ["education"],
            "experience": ["experience"],
            "skills": ["skills"]
        }
        
        for field_info in template["fields"]:
            field_name = field_info["field"]
            
            if field_name in mapping_rules:
                for path in mapping_rules[field_name]:
                    value = self.get_nested_value(extracted_data, path)
                    if value:
                        filled_form["fields"][field_name] = value
                        filled_form["confidence"][field_name] = 0.8
                        break
                
                if field_name not in filled_form["fields"] and field_info["required"]:
                    filled_form["missing_fields"].append(field_name)
        
        return filled_form
    
    def get_nested_value(self, data: Dict, path: str):
        """Lấy giá trị từ nested dictionary"""
        try:
            keys = path.split('.')
            value = data
            for key in keys:
                value = value[key]
            return value
        except (KeyError, TypeError):
            return None
    
    def process_document_file(self, file_path: str, document_type: str = "cv") -> Dict:
        """Xử lý file tài liệu hoàn chỉnh"""
        try:
            # Trích xuất text
            text = self.extract_text_from_file(file_path)
            
            if not text.strip():
                return {"error": "No text could be extracted from the document"}
            
            # Trích xuất thông tin dựa trên loại tài liệu
            if document_type == "cv":
                extracted_data = self.extract_cv_information(text)
            else:
                extracted_data = self.extract_document_information(text, document_type)
            
            return {
                "file_path": file_path,
                "document_type": document_type,
                "extracted_text": text[:500] + "..." if len(text) > 500 else text,
                "extracted_data": extracted_data,
                "processing_status": "success"
            }
            
        except Exception as e:
            return {
                "file_path": file_path,
                "document_type": document_type,
                "error": f"Failed to process document: {e}",
                "processing_status": "failed"
            }
    
    def get_form_templates(self) -> Dict:
        """Lấy danh sách templates biểu mẫu"""
        return self.form_templates
    
    def add_form_template(self, template_name: str, template_data: Dict):
        """Thêm template biểu mẫu mới"""
        self.form_templates[template_name] = template_data
        with open(self.form_templates_file, 'w', encoding='utf-8') as f:
            json.dump(self.form_templates, f, ensure_ascii=False, indent=2)

# Khởi tạo instance global
document_extractor = DocumentExtractor()

